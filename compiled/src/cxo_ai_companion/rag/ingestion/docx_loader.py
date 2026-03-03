"""DOCX (Word) document loader using python-docx."""

from __future__ import annotations

import io
import logging
from dataclasses import dataclass

from cxo_ai_companion.rag.ingestion.base_loader import (
    DocumentLoader,
    DocumentMetadata,
    LoadedDocument,
    LoaderConfig,
)

logger = logging.getLogger(__name__)


@dataclass
class DOCXConfig(LoaderConfig):
    """DOCX-specific loader configuration.

    Attributes:
        extract_tables: Whether to include table content in the output text.
        extract_comments: Whether to extract document comments (not yet implemented).
    """

    extract_tables: bool = True
    extract_comments: bool = False


class DOCXLoader(DocumentLoader):
    """Load text content from DOCX files via python-docx.

    Extracts paragraph text and optionally table content, returning
    a single ``LoadedDocument`` with the combined text.

    Args:
        config: DOCX-specific configuration. Defaults to ``DOCXConfig()``.
    """

    def __init__(self, config: DOCXConfig | None = None) -> None:
        super().__init__(config or DOCXConfig())

    @property
    def docx_config(self) -> DOCXConfig:
        """Return the config cast to ``DOCXConfig``."""
        return self.config  # type: ignore[return-value]

    async def load(self, source: str | bytes) -> list[LoadedDocument]:
        """Load a DOCX from a file path or raw bytes.

        Paragraphs are extracted first, followed by table content (if
        ``extract_tables`` is enabled).  The result is always returned as
        a single :class:`LoadedDocument`.
        """
        import docx  # lazy import — heavy dependency

        if isinstance(source, str):
            try:
                doc = docx.Document(source)
            except Exception as exc:
                logger.error("Failed to open DOCX file %s: %s", source, exc)
                raise
            source_name = source
        else:
            self._validate_size(source)
            try:
                doc = docx.Document(io.BytesIO(source))
            except Exception as exc:
                logger.error("Failed to parse DOCX from bytes: %s", exc)
                raise
            source_name = "bytes_input"

        # ---- Extract paragraphs ----------------------------------------
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

        # ---- Extract tables --------------------------------------------
        table_texts: list[str] = []
        if self.docx_config.extract_tables:
            for table in doc.tables:
                rows: list[str] = []
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells]
                    rows.append(" | ".join(cells))
                table_texts.append("\n".join(rows))

        content_parts = paragraphs + table_texts
        full_text = "\n\n".join(content_parts)

        # ---- Metadata --------------------------------------------------
        metadata = DocumentMetadata(
            word_count=self._count_words(full_text),
        )
        if self.config.extract_metadata and doc.core_properties:
            metadata.title = doc.core_properties.title or ""
            metadata.author = doc.core_properties.author or ""

        logger.debug("Loaded DOCX %s (%d words)", source_name, metadata.word_count)
        return [
            LoadedDocument(
                content=full_text,
                source=source_name,
                mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                metadata=metadata,
            )
        ]

    def supports_source(self, source: str) -> bool:
        """Return ``True`` if *source* ends with ``.docx``."""
        return source.lower().endswith(".docx")


__all__ = [
    "DOCXConfig",
    "DOCXLoader",
]
