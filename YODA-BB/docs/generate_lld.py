"""Generate YODA Low-Level Design (LLD) Document (Word format).

Produces a comprehensive .docx covering architecture, component design,
data flows, API contracts, security, deployment, and operational details.
"""

import sys
sys.path.insert(0, 'foundation/src')

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from datetime import datetime

doc = Document()

# ── Styles ────────────────────────────────────────────────────────
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(10)
style.paragraph_format.space_after = Pt(4)

for level in range(1, 5):
    hs = doc.styles[f'Heading {level}']
    hs.font.name = 'Calibri'
    hs.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)


def add_code(text):
    """Add a code block as a formatted paragraph."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(8)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    return p


def add_table_simple(headers, rows):
    """Add a styled table with headers and data rows."""
    t = doc.add_table(rows=1, cols=len(headers), style='Light Grid Accent 1')
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        t.rows[0].cells[i].text = h
        for run in t.rows[0].cells[i].paragraphs[0].runs:
            run.bold = True
            run.font.size = Pt(9)
    for row_data in rows:
        row = t.add_row()
        for i, val in enumerate(row_data):
            row.cells[i].text = str(val)
            for run in row.cells[i].paragraphs[0].runs:
                run.font.size = Pt(9)
    doc.add_paragraph()
    return t


# ══════════════════════════════════════════════════════════════════
# TITLE PAGE
# ══════════════════════════════════════════════════════════════════
doc.add_paragraph()
doc.add_paragraph()
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('YODA — AI Meeting Companion')
run.font.size = Pt(28)
run.font.color.rgb = RGBColor(0x3B, 0x82, 0xF6)
run.bold = True

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('Low-Level Design Document (LLD)')
run.font.size = Pt(18)
run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

doc.add_paragraph()
meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.add_run(f'Version 1.0 — {datetime.now().strftime("%B %d, %Y")}').font.size = Pt(12)

doc.add_paragraph()
meta2 = doc.add_paragraph()
meta2.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta2.add_run('Confidential — For Internal Use Only').font.italic = True

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
# TOC
# ══════════════════════════════════════════════════════════════════
doc.add_heading('Table of Contents', level=1)
toc_items = [
    '1. Introduction & Scope',
    '2. System Architecture Overview',
    '3. Technology Stack',
    '4. Service Decomposition',
    '5. Foundation Library Design',
    '6. API Design & Contracts',
    '7. Data Model & Database Design',
    '8. RAG Pipeline Design',
    '9. DSPy Integration & Optimization',
    '10. Authentication & Authorization',
    '11. Security Design',
    '12. Real-Time Communication (SSE)',
    '13. Document Classification Engine',
    '14. AI Processing Pipeline',
    '15. Frontend Architecture (Angular)',
    '16. Deployment Architecture',
    '17. Configuration Management',
    '18. Error Handling Strategy',
    '19. Observability & Monitoring',
    '20. Testing Strategy',
    '21. Appendix: API Endpoint Reference',
]
for item in toc_items:
    doc.add_paragraph(item, style='List Number')
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
# 1. INTRODUCTION
# ══════════════════════════════════════════════════════════════════
doc.add_heading('1. Introduction & Scope', level=1)

doc.add_heading('1.1 Purpose', level=2)
doc.add_paragraph(
    'This Low-Level Design (LLD) document describes the detailed technical design of YODA '
    '(Your Organizational Digital Assistant), an AI-powered meeting companion for enterprise CXOs. '
    'It covers component-level architecture, data flows, API contracts, security mechanisms, '
    'and deployment topology at a level sufficient for developers to implement, maintain, and extend the system.'
)

doc.add_heading('1.2 Scope', level=2)
doc.add_paragraph('This document covers:')
for item in [
    '6 backend microservices (Python/FastAPI)',
    '1 shared foundation library (yoda_foundation)',
    '1 Angular 21 frontend (yoda-frontend)',
    'PostgreSQL with pgvector for RAG',
    'Redis for caching',
    'Nginx API gateway',
    'Docker Compose deployment',
    'Microsoft 365 integration (Graph API, ACS, AI Foundry, Entra ID)',
]:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('1.3 Audience', level=2)
doc.add_paragraph('Backend developers, frontend developers, DevOps engineers, security reviewers, and QA engineers.')

doc.add_heading('1.4 References', level=2)
add_table_simple(
    ['Document', 'Location'],
    [
        ('User Stories & Requirements', 'docs/YODA_User_Stories_Requirements.docx'),
        ('Database Schema Reference', 'docs/YODA_Database_Schema_Reference.docx'),
        ('Wireframe Prototype', 'Notes/CXO_AI_Companion_Wireframes_v2.html'),
        ('Feature Mapping CSV', 'docs/wireframe-backend-mapping.csv'),
        ('Golden QA Cases', 'config/golden_qa.json'),
        ('CI/CD Pipeline', '.github/workflows/ci.yml'),
    ],
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
# 2. SYSTEM ARCHITECTURE
# ══════════════════════════════════════════════════════════════════
doc.add_heading('2. System Architecture Overview', level=1)

doc.add_heading('2.1 Architecture Pattern', level=2)
doc.add_paragraph(
    'YODA follows a microservices architecture with a shared foundation library. '
    'All services are independently deployable Docker containers, communicating through a shared '
    'PostgreSQL database and an Nginx reverse proxy (API gateway). '
    'The frontend is a single-page application (SPA) served by its own Nginx container.'
)

doc.add_heading('2.2 Component Diagram', level=2)
add_code(
    'Browser (Angular SPA)\n'
    '    │\n'
    '    ▼\n'
    'Nginx Gateway (:80)\n'
    '    ├── /api/meetings, /api/action-items  →  meeting-service (:8010)\n'
    '    ├── /api/documents                    →  document-service (:8011)\n'
    '    ├── /api/chat                         →  chat-service (:8012)\n'
    '    ├── /api/dashboard, /api/insights     →  dashboard-service (:8013)\n'
    '    ├── /api/briefs                       →  pre-meeting-brief (:8014)\n'
    '    ├── /api/digests                      →  weekly-digest (:8015)\n'
    '    └── /* (catch-all)                    →  frontend (:4200)\n'
    '                    │\n'
    '    ┌───────────────┴───────────────┐\n'
    '    ▼                               ▼\n'
    'PostgreSQL + pgvector           Redis (cache)\n'
    '                                    │\n'
    '                           Azure Services\n'
    '                           ├── Graph API\n'
    '                           ├── AI Foundry (GPT-4o)\n'
    '                           ├── ACS (Call Automation)\n'
    '                           └── Entra ID (Auth)'
)

doc.add_heading('2.3 Design Principles', level=2)
for principle, desc in [
    ('Async-first', 'All I/O uses async/await with asyncpg, httpx, and asyncio'),
    ('Session-per-operation', 'Each DB operation creates its own session from a factory — no shared long-lived sessions'),
    ('Config from environment', 'All settings via pydantic-settings, never hardcoded'),
    ('Lazy initialization', 'Services and singletons created in lifespan, stored on app.state'),
    ('Tenant isolation', 'All queries scoped by authenticated user_id'),
    ('Fail-safe', 'Optional components (Graph, AI, Redis) degrade gracefully if unavailable'),
]:
    p = doc.add_paragraph()
    p.add_run(f'{principle}: ').bold = True
    p.add_run(desc)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
# 3. TECHNOLOGY STACK
# ══════════════════════════════════════════════════════════════════
doc.add_heading('3. Technology Stack', level=1)

add_table_simple(
    ['Layer', 'Technology', 'Version', 'Purpose'],
    [
        ('Backend Framework', 'FastAPI', '0.100+', 'Async REST API framework'),
        ('Language', 'Python', '3.11+', 'Backend services'),
        ('ORM', 'SQLAlchemy', '2.0 (async)', 'Database access with mapped_column'),
        ('Database', 'PostgreSQL', '16', 'Primary data store'),
        ('Vector Search', 'pgvector', '0.5+', '1536-dim cosine similarity search'),
        ('Cache', 'Redis', '7', 'Session cache, rate limiting'),
        ('Task Scheduler', 'APScheduler', '3.10+', 'Nudges, digest generation, subscription renewal'),
        ('AI/LLM', 'Azure AI Foundry', 'GPT-4o / 4o-mini', 'Summary extraction, Q&A, classification'),
        ('Embeddings', 'Azure OpenAI', 'text-embedding-3-small', '1536-dim document embeddings'),
        ('DSPy', 'dspy-ai', '3.1.3', 'Prompt optimization with BootstrapFewShot'),
        ('Auth', 'MSAL / Entra ID', 'v2.0', 'OAuth2 JWT token validation'),
        ('Microsoft APIs', 'Graph API', 'v1.0', 'Calendar, users, email, SharePoint, OneDrive'),
        ('Call Platform', 'Azure Communication Services', '-', 'Bot join/leave meeting calls'),
        ('Frontend Framework', 'Angular', '21', 'Standalone components, signals'),
        ('Frontend Auth', 'MSAL.js', '3.x', 'Azure AD browser authentication'),
        ('API Gateway', 'Nginx', 'alpine', 'Path-based reverse proxy'),
        ('Containers', 'Docker Compose', '3.9', '10 containers'),
        ('CI/CD', 'GitHub Actions', '-', 'Matrix test jobs, lint, build'),
        ('Migrations', 'Alembic', '1.13+', 'Async PostgreSQL migrations'),
    ],
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
# 4. SERVICE DECOMPOSITION
# ══════════════════════════════════════════════════════════════════
doc.add_heading('4. Service Decomposition', level=1)

services = [
    ('meeting-service', '8010', '20', '52',
     'Meeting lifecycle, transcription, AI processing, action items, nudges, calendar webhooks, bot control',
     ['routes/meetings.py — CRUD + join/leave/reprocess',
      'routes/action_items.py — CRUD + complete/snooze + owner filters',
      'routes/bot_events.py — transcript chunks + lifecycle events',
      'routes/webhooks.py — Graph calendar notifications (with clientState validation)',
      'routes/sse.py — Server-Sent Events for live transcript streaming',
      'routes/admin.py — User management (admin only)',
      'services/ai_processor.py — GPT-4o transcript extraction',
      'services/bot_commander.py — .NET media bot HTTP control',
      'services/calendar_watcher.py — Graph subscription lifecycle',
      'services/delivery.py — Adaptive Card delivery to Teams',
      'services/nudge_scheduler.py — APScheduler background nudges',
      'services/owner_resolver.py — Fuzzy name → Graph user resolution',
      'services/post_processing.py — Orchestrates AI → resolve → deliver',
      'services/graph_client.py — Microsoft Graph API wrapper',
      'services/meeting_tag_service.py — Computes meeting tags',
      ]),
    ('document-service', '8011', '12', '49',
     'Document management, SharePoint sync, RAG ingestion, semantic search, classification, email indexing',
     ['routes/documents.py — CRUD + sync + search + classify + index-emails',
      'services/document_service.py — Sync, process, search, classify, email indexing',
      'services/email_indexer.py — Fetch Graph emails, chunk, embed into pgvector',
      'dependencies.py — RAG singletons (embedder, vector store, chunker, classifier)',
      ]),
    ('chat-service', '8012', '8', '13',
     'RAG-powered AI Q&A with citation tracking',
     ['routes/chat.py — Session CRUD + send message',
      'services/chat_service.py — RAG query → DSPy ChainOfThought → citations',
      ]),
    ('dashboard-service', '8013', '16', '46',
     'Executive dashboard, insights, notifications, search, recommendations, conflict detection',
     ['routes/dashboard.py — Stats, attention items, activity feed, recommendations',
      'routes/insights.py — Meeting time, completion rate, velocity, collaboration, recurring topics',
      'routes/notifications.py — CRUD + mark read',
      'routes/search.py — Cross-entity search',
      'routes/user_settings.py — GET/PATCH user preferences',
      'routes/health.py — Health check + M365 status',
      'services/dashboard_service.py — Aggregate queries',
      'services/insight_service.py — Analytics computations',
      'services/notification_service.py — Notification management',
      'services/conflict_detection_service.py — Decision contradiction detection',
      'services/topic_detection_service.py — Recurring topic analysis',
      ]),
    ('pre-meeting-brief-service', '8014', '8', '10',
     'AI-generated pre-meeting preparation briefs',
     ['routes/briefs.py — GET /api/briefs/{meeting_id}',
      'services/pre_meeting_service.py — Concurrent: attendees, decisions, docs, emails, AI questions',
      'services/graph_client.py — Graph API for attendee context',
      ]),
    ('weekly-digest-service', '8015', '8', '10',
     'Auto-generated weekly summaries via APScheduler',
     ['routes/digests.py — GET latest + POST generate',
      'services/weekly_digest_service.py — Aggregate week, AI summarize, deliver',
      ]),
]

for name, port, files, tests, desc, components in services:
    doc.add_heading(f'4.x {name} (:{port})', level=2)
    p = doc.add_paragraph()
    p.add_run('Purpose: ').bold = True
    p.add_run(desc)
    p = doc.add_paragraph()
    p.add_run(f'Files: {files} | Tests: {tests}')

    doc.add_paragraph('Key Components:', style='List Bullet')
    for comp in components:
        doc.add_paragraph(comp, style='List Bullet 2')
    doc.add_paragraph()

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
# 5. FOUNDATION LIBRARY
# ══════════════════════════════════════════════════════════════════
doc.add_heading('5. Foundation Library Design', level=1)
doc.add_paragraph(
    'The yoda_foundation library (241 Python files) is a shared package installed as an editable '
    'dependency by all 6 services. It provides models, schemas, RAG pipeline, DSPy integration, '
    'security, resilience, observability, guardrails, events, and memory management.'
)

add_table_simple(
    ['Module', 'Files', 'Purpose'],
    [
        ('models/', '12', 'SQLAlchemy 2.0 async ORM models with UUID PKs + timestamps'),
        ('schemas/', '15', 'Pydantic v2 request/response schemas (from_attributes=True)'),
        ('config/', '2', 'Settings via pydantic-settings (90+ env vars)'),
        ('exceptions/', '10', 'YodaBaseException hierarchy with severity/category/retryability'),
        ('security/', '20', 'JWT validation, RBAC (3 roles), data governance, secrets management'),
        ('data_access/', '20', 'BaseConnector, connection pools, Graph/ACS/AI Foundry connectors, repositories'),
        ('rag/', '28+', 'Chunking, embeddings, vector store, retrieval (hybrid, reranker, HyDE), classification, evaluation'),
        ('dspy/', '8+', 'Custom modules + real dspy-ai integration, signatures, optimizer'),
        ('resilience/', '32', 'Retry, circuit breaker, bulkhead, fallback, timeout, dead letter, recovery'),
        ('observability/', '10', 'OpenTelemetry tracing, metrics, structured JSON logging'),
        ('guardrails/', '12', 'Content safety, jailbreak detection, fact-checking, moderation'),
        ('events/', '25', 'Event bus (in-memory + Redis), handlers, triggers, streaming, sourcing'),
        ('memory/', '16', 'Tiered memory (working, episodic, semantic, procedural), consolidation, decay'),
        ('middleware/', '6', 'Correlation ID, error handler, rate limiter, security headers'),
        ('utils/', '9', 'MSAL token provider, caching (memory + Redis), retry decorator'),
        ('templates/', '5', 'Adaptive Card JSON templates for Teams delivery'),
    ],
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
# 6. API DESIGN
# ══════════════════════════════════════════════════════════════════
doc.add_heading('6. API Design & Contracts', level=1)

doc.add_heading('6.1 Design Conventions', level=2)
for conv, desc in [
    ('URL Pattern', '/api/{resource} for collections, /api/{resource}/{id} for items'),
    ('Methods', 'GET (read), POST (create/action), PATCH (partial update), DELETE (remove)'),
    ('Auth', 'Bearer JWT in Authorization header (MSAL → Entra ID)'),
    ('Response Format', 'JSON with snake_case field names'),
    ('Pagination', 'limit/offset query params, response includes total count'),
    ('Errors', '400 (validation), 401 (unauthorized), 403 (forbidden), 404 (not found), 429 (rate limit), 502 (upstream)'),
    ('Tenant Scope', 'All list endpoints filter by authenticated user_id'),
]:
    p = doc.add_paragraph()
    p.add_run(f'{conv}: ').bold = True
    p.add_run(desc)

doc.add_heading('6.2 Endpoint Summary', level=2)

endpoints = [
    # Meeting service
    ('GET', '/api/meetings', 'meeting', 'List meetings (with computed tags)'),
    ('GET', '/api/meetings/{id}', 'meeting', 'Get meeting detail (summary + actions + participants)'),
    ('POST', '/api/meetings', 'meeting', 'Create meeting'),
    ('POST', '/api/meetings/{id}/join', 'meeting', 'Trigger bot to join meeting'),
    ('POST', '/api/meetings/{id}/leave', 'meeting', 'Remove bot from meeting'),
    ('POST', '/api/meetings/{id}/reprocess', 'meeting', 'Re-run AI processing'),
    ('GET', '/api/meetings/{id}/transcript', 'meeting', 'Get transcript segments'),
    ('GET', '/api/meetings/{id}/events', 'meeting', 'SSE stream of live updates'),
    ('PATCH', '/api/meetings/{id}/summary', 'meeting', 'Edit AI summary'),
    ('GET', '/health/browser-bot', 'meeting', 'Bot health check'),
    ('GET', '/api/action-items', 'meeting', 'List action items (filterable)'),
    ('PATCH', '/api/action-items/{id}', 'meeting', 'Update action item'),
    ('POST', '/api/action-items/{id}/complete', 'meeting', 'Mark complete'),
    ('POST', '/api/action-items/{id}/snooze', 'meeting', 'Snooze nudges'),
    ('GET', '/api/admin/users', 'meeting', 'List users (admin)'),
    ('POST', '/api/admin/users', 'meeting', 'Create user (admin)'),
    ('PATCH', '/api/admin/users/{id}', 'meeting', 'Update user (admin)'),
    ('DELETE', '/api/admin/users/{id}', 'meeting', 'Delete user (admin)'),
    ('POST', '/api/bot-events/transcript', 'meeting', 'Ingest transcript chunks'),
    ('POST', '/api/bot-events/lifecycle', 'meeting', 'Bot lifecycle events'),
    ('POST', '/webhooks/graph', 'meeting', 'Graph webhook receiver'),
    # Document service
    ('GET', '/api/documents', 'document', 'List documents (type filter, sort)'),
    ('GET', '/api/documents/{id}', 'document', 'Get single document'),
    ('GET', '/api/documents/search', 'document', 'Semantic search (RAG)'),
    ('GET', '/api/documents/recent', 'document', 'Recently updated'),
    ('GET', '/api/documents/needs-review', 'document', 'Docs needing review'),
    ('GET', '/api/documents/shared-with-me', 'document', 'Shared via Graph'),
    ('GET', '/api/documents/meeting-related', 'document', 'Docs for today\'s meetings'),
    ('POST', '/api/documents/sync', 'document', 'Sync from SharePoint/OneDrive'),
    ('POST', '/api/documents/upload', 'document', 'Upload + RAG ingestion'),
    ('POST', '/api/documents/{id}/classify', 'document', 'AI classify + persist'),
    ('POST', '/api/documents/classify-text', 'document', 'Classify arbitrary text'),
    ('POST', '/api/documents/{id}/reprocess', 'document', 'Re-index via pipeline'),
    ('POST', '/api/documents/index-emails', 'document', 'Index emails from Graph'),
    # Chat service
    ('POST', '/api/chat/sessions', 'chat', 'Create chat session'),
    ('GET', '/api/chat/sessions', 'chat', 'List sessions'),
    ('POST', '/api/chat/sessions/{id}/messages', 'chat', 'Send message → RAG AI response'),
    ('GET', '/api/chat/sessions/{id}/messages', 'chat', 'Get message history'),
    # Dashboard service
    ('GET', '/api/dashboard/stats', 'dashboard', 'KPI summary'),
    ('GET', '/api/dashboard/attention-items', 'dashboard', 'Items needing attention'),
    ('GET', '/api/dashboard/activity-feed', 'dashboard', 'Recent activity'),
    ('GET', '/api/dashboard/recommendations', 'dashboard', 'AI recommendations'),
    ('GET', '/api/dashboard/m365-status', 'dashboard', 'M365 connection check'),
    ('GET', '/api/insights/meeting-time', 'dashboard', 'Meeting time analysis'),
    ('GET', '/api/insights/action-completion', 'dashboard', 'Completion rate'),
    ('GET', '/api/insights/decision-velocity', 'dashboard', 'Decision resolution speed'),
    ('GET', '/api/insights/collaboration', 'dashboard', 'Collaboration patterns'),
    ('GET', '/api/insights/patterns', 'dashboard', 'Notable patterns'),
    ('GET', '/api/insights/recurring-topics', 'dashboard', 'Topics in 3+ meetings'),
    ('GET', '/api/notifications', 'dashboard', 'List notifications'),
    ('GET', '/api/notifications/count', 'dashboard', 'Unread count'),
    ('PATCH', '/api/notifications/{id}/read', 'dashboard', 'Mark read'),
    ('POST', '/api/notifications/read-all', 'dashboard', 'Mark all read'),
    ('GET', '/api/search', 'dashboard', 'Cross-entity search'),
    ('GET', '/api/settings', 'dashboard', 'Get user preferences'),
    ('PATCH', '/api/settings', 'dashboard', 'Update user preferences'),
    # Brief service
    ('GET', '/api/briefs/{meeting_id}', 'brief', 'Generate pre-meeting brief'),
    # Digest service
    ('GET', '/api/digests/latest', 'digest', 'Get latest digest'),
    ('POST', '/api/digests/generate', 'digest', 'Trigger digest generation'),
    # Health (all services)
    ('GET', '/health', 'all', 'Service health check'),
]

add_table_simple(
    ['Method', 'Path', 'Service', 'Description'],
    [(e[0], e[1], e[2], e[3]) for e in endpoints],
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
# 7. DATA MODEL
# ══════════════════════════════════════════════════════════════════
doc.add_heading('7. Data Model & Database Design', level=1)
doc.add_paragraph(
    'See separate document: docs/YODA_Database_Schema_Reference.docx for complete table definitions, '
    'columns, relationships, indexes, and migration commands.'
)

doc.add_heading('7.1 ER Summary', level=2)
add_code(
    'meetings (1) ──┬── (N) meeting_participants\n'
    '               ├── (N) transcript_segments\n'
    '               ├── (1) meeting_summaries\n'
    '               ├── (N) action_items\n'
    '               ├── (N) documents\n'
    '               ├── (N) meeting_insights\n'
    '               └── (M:N) projects (via project_meetings)\n'
    '\n'
    'documents (1) ── (N) document_chunks (with VECTOR(1536) embedding)\n'
    'chat_sessions (1) ── (N) chat_messages'
)

doc.add_heading('7.2 Key Design Decisions', level=2)
for decision, rationale in [
    ('UUID primary keys', 'Globally unique, no sequence conflicts across services, safe for distributed systems'),
    ('Single shared database', 'Simplicity for v1.0 — all services share one PostgreSQL. Schema changes require coordinated deploys.'),
    ('pgvector for embeddings', 'Native PostgreSQL extension — no separate vector DB needed. Cosine distance operator (<=>).'),
    ('JSON columns for flexible data', 'decisions, key_topics, sources stored as JSON — schema-less for evolving AI output format'),
    ('lazy="raise" on all relationships', 'Prevents accidental synchronous queries in async context (MissingGreenlet). Forces explicit selectinload().'),
    ('Frozen PGVectorConfig', 'Prevents SQL injection via table name mutation after dataclass validation'),
]:
    p = doc.add_paragraph()
    p.add_run(f'{decision}: ').bold = True
    p.add_run(rationale)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
# 8. RAG PIPELINE
# ══════════════════════════════════════════════════════════════════
doc.add_heading('8. RAG Pipeline Design', level=1)

doc.add_heading('8.1 Ingestion Pipeline', level=2)
add_code(
    'Document (PDF/DOCX/PPTX/HTML/CSV/Email)\n'
    '    │\n'
    '    ▼ Format Loader (extracts raw text)\n'
    '    │\n'
    '    ▼ RecursiveChunker (2048 chars / ~512 tokens, sentence boundary separators)\n'
    '    │\n'
    '    ▼ AzureEmbedder (text-embedding-3-small → 1536-dim vectors)\n'
    '    │\n'
    '    ▼ PGVectorStore (INSERT INTO document_chunks with ON CONFLICT upsert)\n'
    '    │\n'
    '    ▼ DocumentClassifier (23 categories, keyword boost, filename hints)'
)

doc.add_heading('8.2 Query Pipeline', level=2)
add_code(
    'User Question\n'
    '    │\n'
    '    ▼ QueryExpander (HyDE: generate hypothetical answer, embed both, average)\n'
    '    │\n'
    '    ▼ HybridRetriever\n'
    '    │   ├── Vector search (pgvector cosine, top-15)\n'
    '    │   └── Full-text search (PostgreSQL tsvector/tsquery, top-15)\n'
    '    │   └── Reciprocal Rank Fusion merge\n'
    '    │\n'
    '    ▼ MetadataFilter (auto-detect time/type/meeting from query)\n'
    '    │\n'
    '    ▼ LLMReranker (GPT-4o scores relevance 1-10, returns top-5)\n'
    '    │\n'
    '    ▼ ContextBuilder (assemble chunks with [N] markers)\n'
    '    │\n'
    '    ▼ CitationTracker (register sources)\n'
    '    │\n'
    '    ▼ DSPy ChainOfThought (ContextualQA signature)\n'
    '    │   ├── dspy.Suggest: answer must cite sources\n'
    '    │   └── dspy.Suggest: answer must be >20 chars\n'
    '    │\n'
    '    ▼ CitationTracker.resolve_citations(answer)\n'
    '    │\n'
    '    ▼ RAGResult { answer, sources, citations, confidence, rationale }'
)

doc.add_heading('8.3 Evaluation Framework', level=2)
doc.add_paragraph('Golden QA cases stored in config/golden_qa.json (20 cases, 8 categories):')
add_table_simple(
    ['Category', 'Count', 'Tests'],
    [
        ('Factual', '6', 'Exact data retrieval from single documents'),
        ('Reasoning', '3', 'Why/how questions requiring inference'),
        ('Multi-doc', '3', 'Answers requiring information from 2+ documents'),
        ('Meeting', '2', 'Meeting-specific context questions'),
        ('Policy', '1', 'Organizational policy questions'),
        ('Compliance', '1', 'Legal/compliance document questions'),
        ('Temporal', '2', 'Time-based trend questions'),
        ('Comparison', '2', 'Comparing entities or metrics'),
    ],
)

doc.add_paragraph('Metrics: precision@k, recall@k, MRR, LLM-judged answer relevance, faithfulness scoring.')

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
# 9. DSPy
# ══════════════════════════════════════════════════════════════════
doc.add_heading('9. DSPy Integration & Optimization', level=1)

doc.add_heading('9.1 Architecture', level=2)
doc.add_paragraph(
    'YODA uses the real dspy-ai library (v3.1.3) alongside a backward-compatible custom framework. '
    'The integration layer in dspy/integration.py bridges Azure AI Foundry with DSPy\'s optimization pipeline.'
)

add_table_simple(
    ['Component', 'Class', 'Purpose'],
    [
        ('configure_dspy()', 'Function', 'One-call Azure AI Foundry configuration'),
        ('ContextualQA', 'dspy.Signature', 'RAG Q&A: contexts + question → reasoning + answer + confidence + citations'),
        ('MeetingExtraction', 'dspy.Signature', 'Transcript → summary + action_items + decisions + topics + questions'),
        ('DocumentSummary', 'dspy.Signature', 'Document text → summary + key_points + entities'),
        ('InsightDetection', 'dspy.Signature', 'Current vs past decisions → conflicts + severity + recommendation'),
        ('YodaQA', 'dspy.Module', 'ChainOfThought(ContextualQA) + Suggest assertions'),
        ('YodaMeetingExtractor', 'dspy.Module', 'ChainOfThought(MeetingExtraction) + validation'),
        ('YodaInsightDetector', 'dspy.Module', 'ChainOfThought(InsightDetection) + severity validation'),
        ('YodaOptimizer', 'Class', 'BootstrapFewShot optimization from JSON golden QA cases'),
    ],
)

doc.add_heading('9.2 Optimization Flow', level=2)
add_code(
    'config/golden_qa.json (20 cases)\n'
    '    │\n'
    '    ▼ YodaOptimizer.load_from_json()\n'
    '    │\n'
    '    ▼ dspy.BootstrapFewShot(metric=default_metric)\n'
    '    │\n'
    '    ▼ Compile: runs each case, selects best few-shot demos\n'
    '    │\n'
    '    ▼ Optimized YodaQA module (auto-tuned prompts)'
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
# 10. AUTH
# ══════════════════════════════════════════════════════════════════
doc.add_heading('10. Authentication & Authorization', level=1)

doc.add_heading('10.1 Auth Flow', level=2)
add_code(
    'Browser → MSAL.js → Azure AD login\n'
    '    ↓ JWT access token\n'
    'Angular HTTP Interceptor → Authorization: Bearer <token>\n'
    '    ↓\n'
    'Backend get_current_user() dependency\n'
    '    ↓ Validates signature (RS256, JWKS from Microsoft)\n'
    '    ↓ Extracts claims: oid, tid, name, preferred_username, roles\n'
    '    ↓ Maps roles → permissions via _ROLE_PERMISSION_MAP\n'
    '    ↓\n'
    'SecurityContext { user_id, tenant_id, permissions, roles, metadata }'
)

doc.add_heading('10.2 RBAC Roles', level=2)
add_table_simple(
    ['Role', 'Permissions'],
    [
        ('CXO.Admin', 'meetings.*, documents.*, insights.*, notifications.*, projects.*, search.*, admin.*'),
        ('CXO.User', 'meetings.read/write, documents.read/write, insights.read, notifications.read/write, search.read'),
        ('CXO.Viewer', 'meetings.read, documents.read, insights.read, notifications.read, search.read'),
    ],
)

doc.add_heading('10.3 Dev Mode', level=2)
doc.add_paragraph(
    'When REQUIRE_AUTH=false, requests without tokens receive an anonymous SecurityContext. '
    'Frontend skips MSAL when environment.requireAuth=false.'
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
# 11. SECURITY
# ══════════════════════════════════════════════════════════════════
doc.add_heading('11. Security Design', level=1)

add_table_simple(
    ['Threat', 'Mitigation', 'Implementation'],
    [
        ('Webhook forgery', 'clientState validation', 'webhooks.py validates against expected value'),
        ('SQL injection', 'ORM parameterized + frozen dataclass', 'PGVectorConfig(frozen=True)'),
        ('XSS', 'Angular sanitization', 'Zero innerHTML, zero bypassSecurityTrust'),
        ('CSRF', 'Bearer token auth', 'No cookie-based auth'),
        ('Event loop blocking', 'asyncio.to_thread()', 'JWKS fetch wrapped in thread'),
        ('Session corruption', 'Session-per-operation', 'session_factory pattern, no shared sessions'),
        ('Data leakage (multi-tenant)', 'Tenant-scoped queries', 'All list endpoints filter by user_id'),
        ('Rate limiting', 'In-memory limiter', '100 req/min/IP on webhook endpoint'),
        ('Error info leak', 'Generic error messages', 'Raw exceptions logged server-side only'),
        ('CORS', 'Restricted methods/headers', 'Only GET/POST/PATCH/DELETE/OPTIONS + specific headers'),
        ('Lazy loading crash', 'lazy="raise"', 'All model relationships require explicit eager loading'),
        ('Auth bypass', 'Startup CRITICAL log', 'Warns when AZURE_AD_AUDIENCE empty + REQUIRE_AUTH=true'),
    ],
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
# 12. SSE
# ══════════════════════════════════════════════════════════════════
doc.add_heading('12. Real-Time Communication (SSE)', level=1)

add_code(
    '.NET Media Bot → POST /api/bot-events/transcript\n'
    '    │\n'
    '    ▼ bot_events.py saves segments to DB\n'
    '    ▼ publish_meeting_event(meeting_id, {type: "transcript_update", ...})\n'
    '    │\n'
    '    ▼ In-memory event queues (per-meeting, per-client)\n'
    '    │\n'
    '    ▼ GET /api/meetings/{id}/events (SSE stream)\n'
    '    ▼ StreamingResponse(text/event-stream)\n'
    '    ▼ Keepalive heartbeat every 30 seconds\n'
    '    ▼ Auto-cleanup on client disconnect'
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
# 13. CLASSIFICATION
# ══════════════════════════════════════════════════════════════════
doc.add_heading('13. Document Classification Engine', level=1)

doc.add_paragraph('23 enterprise document categories with vector template matching:')
add_table_simple(
    ['Category', 'Examples', 'Default Priority'],
    [
        ('MBR', 'Monthly Business Review', 'Medium'),
        ('QBR', 'Quarterly Business Review', 'Medium'),
        ('SOW', 'Statement of Work', 'Medium'),
        ('MSA', 'Master Service Agreement', 'High'),
        ('Status Report', 'Weekly/Sprint status', 'Medium'),
        ('Delivery Document', 'Deliverables, handoffs', 'Medium'),
        ('Risk Document', 'Risk registers, mitigation', 'High'),
        ('MOM', 'Minutes of Meeting', 'Medium'),
        ('Escalation', 'Escalation notices', 'High'),
        ('Contract / NDA', 'Legal agreements', 'High'),
        ('Invoice / Financial', 'Billing, forecasts', 'Medium'),
        ('Technical Spec', 'Architecture, design docs', 'Low'),
        ('Presentation', 'Decks, pitch materials', 'Low'),
        ('Resume / HR', 'Hiring documents', 'Low'),
        ('General', 'Catch-all category', 'Low'),
    ],
)

doc.add_paragraph(
    'Classification uses: cosine similarity against 100+ pre-embedded templates, '
    'keyword boosting (per-category strong/medium keywords), '
    'filename pattern matching, and MIME type signals.'
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
# 14. AI PROCESSING
# ══════════════════════════════════════════════════════════════════
doc.add_heading('14. AI Processing Pipeline', level=1)

doc.add_heading('14.1 Meeting Transcript Processing', level=2)
add_code(
    'Meeting ends (bot_lifecycle event)\n'
    '    │\n'
    '    ▼ PostProcessingService.run(meeting_id)\n'
    '    │   creates fresh DB session\n'
    '    │\n'
    '    ▼ Load transcript segments (ordered by sequence_number)\n'
    '    │\n'
    '    ▼ AIProcessor.process_meeting()\n'
    '    │   Short meetings (<30 min): single GPT-4o-mini call\n'
    '    │   Long meetings (≥30 min): chunked processing with GPT-4o\n'
    '    │\n'
    '    ▼ Extract: summary, action_items (with confidence), decisions, key_topics, unresolved_questions\n'
    '    │\n'
    '    ▼ OwnerResolver: fuzzy name → Graph user ID (rapidfuzz matching)\n'
    '    │\n'
    '    ▼ DeliveryService: send Adaptive Card to Teams chat\n'
    '    │\n'
    '    ▼ ConflictDetectionService: compare decisions against past 90 days'
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
# 15. FRONTEND
# ══════════════════════════════════════════════════════════════════
doc.add_heading('15. Frontend Architecture (Angular 21)', level=1)

doc.add_heading('15.1 Project Structure', level=2)
add_code(
    'yoda-frontend/src/app/\n'
    '├── auth/              MSAL configuration\n'
    '├── core/\n'
    '│   ├── models/        TypeScript interfaces (mirror Pydantic schemas)\n'
    '│   └── services/      HTTP services (one per backend service)\n'
    '├── layout/\n'
    '│   ├── shell/         Flex container (sidebar + topbar + router-outlet)\n'
    '│   ├── sidebar/       7 nav items + settings\n'
    '│   └── topbar/        Search dropdown, notifications, avatar\n'
    '├── features/\n'
    '│   ├── dashboard/     4 stat cards, meetings, attention, activity, shortcuts\n'
    '│   ├── meetings/      List (grouped by day) + detail (summary)\n'
    '│   ├── brief/         Pre-meeting brief (attendees, decisions, docs, questions)\n'
    '│   ├── action-items/  4 urgency sections, filters, complete/snooze\n'
    '│   ├── chat/          AI Q&A with citations\n'
    '│   ├── documents/     Type filter, search, review, sync\n'
    '│   ├── insights/      4 insight cards + patterns\n'
    '│   ├── digest/        Weekly summary with stats\n'
    '│   ├── settings/      Theme, notification prefs\n'
    '│   └── not-found/     404 page\n'
    '└── shared/\n'
    '    └── utils/         formatRelative, formatTime, getInitials, computeDuration'
)

doc.add_heading('15.2 Key Patterns', level=2)
for pattern, desc in [
    ('Standalone components', 'Every component is standalone: true — no NgModules'),
    ('Signals', 'All state managed via signal() and computed()'),
    ('OnPush change detection', 'All components use ChangeDetectionStrategy.OnPush'),
    ('takeUntilDestroyed', 'All subscriptions cleaned up via DestroyRef'),
    ('Lazy-loaded routes', 'Every feature uses loadComponent() for code splitting'),
    ('Conditional MSAL', 'Auth providers included only when environment.requireAuth=true'),
    ('Dynamic user identity', 'UserService.profile signal populated from MSAL context'),
    ('ARIA accessibility', 'role, aria-label, aria-expanded on interactive elements'),
    ('CSS custom properties', 'Dark/light theme via --bg-primary, --text-primary, etc.'),
    ('Proxy config', 'proxy.conf.json maps /api/* to backend ports (8010-8015)'),
]:
    p = doc.add_paragraph()
    p.add_run(f'{pattern}: ').bold = True
    p.add_run(desc)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
# 16. DEPLOYMENT
# ══════════════════════════════════════════════════════════════════
doc.add_heading('16. Deployment Architecture', level=1)

doc.add_heading('16.1 Docker Compose Topology', level=2)
add_table_simple(
    ['Container', 'Image', 'Port', 'Depends On'],
    [
        ('postgres', 'pgvector/pgvector:pg16', '5432', '—'),
        ('redis', 'redis:7-alpine', '6379', '—'),
        ('meeting-service', 'Custom Dockerfile', '8010', 'postgres, redis'),
        ('document-service', 'Custom Dockerfile', '8011', 'postgres, redis'),
        ('chat-service', 'Custom Dockerfile', '8012', 'postgres, redis'),
        ('dashboard-service', 'Custom Dockerfile', '8013', 'postgres, redis'),
        ('pre-meeting-brief', 'Custom Dockerfile', '8014', 'postgres, redis'),
        ('weekly-digest', 'Custom Dockerfile', '8015', 'postgres, redis'),
        ('frontend', 'node:20 → nginx:alpine', '4200', '—'),
        ('nginx', 'nginx:alpine', '80', 'All services + frontend'),
    ],
)

doc.add_heading('16.2 Dockerfile Pattern (Backend)', level=2)
add_code(
    'FROM python:3.11-slim AS base\n'
    'COPY foundation/ → pip install\n'
    'COPY services/{name}/ → pip install\n'
    'EXPOSE {port}\n'
    'CMD ["uvicorn", "{name}_service.main:app", "--host", "0.0.0.0", "--port", "{port}"]'
)

doc.add_heading('16.3 Dockerfile Pattern (Frontend)', level=2)
add_code(
    'FROM node:20-alpine AS builder\n'
    'RUN npm ci && npx ng build --configuration=production\n'
    '\n'
    'FROM nginx:alpine\n'
    'COPY dist/ → /usr/share/nginx/html\n'
    'try_files $uri $uri/ /index.html (SPA routing)'
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
# 17. CONFIGURATION
# ══════════════════════════════════════════════════════════════════
doc.add_heading('17. Configuration Management', level=1)

doc.add_paragraph('All configuration via environment variables (pydantic-settings). Key variables:')

add_table_simple(
    ['Variable', 'Required', 'Default', 'Purpose'],
    [
        ('DATABASE_URL', 'Yes', '—', 'PostgreSQL asyncpg connection string'),
        ('REDIS_URL', 'No', '—', 'Redis connection (optional caching)'),
        ('AZURE_TENANT_ID', 'Yes*', '—', 'Entra ID tenant'),
        ('AZURE_CLIENT_ID', 'Yes*', '—', 'App registration client ID'),
        ('AZURE_CLIENT_SECRET', 'Yes*', '—', 'App registration secret'),
        ('AI_FOUNDRY_ENDPOINT', 'No', '—', 'Azure AI Foundry URL'),
        ('AI_FOUNDRY_API_KEY', 'No', '—', 'AI Foundry API key'),
        ('REQUIRE_AUTH', 'No', 'true', 'Set false for dev mode'),
        ('CORS_ALLOWED_ORIGINS', 'No', '[]', 'Allowed CORS origins'),
        ('PORT', 'No', 'varies', 'Service listen port'),
        ('DEBUG', 'No', 'false', 'Enable debug logging'),
    ],
)
doc.add_paragraph('* Required for production. Can be empty for local dev with REQUIRE_AUTH=false.')

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
# 18. ERROR HANDLING
# ══════════════════════════════════════════════════════════════════
doc.add_heading('18. Error Handling Strategy', level=1)

doc.add_heading('18.1 Exception Hierarchy', level=2)
add_code(
    'YodaBaseException\n'
    '├── AuthenticationError\n'
    '├── AuthorizationError\n'
    '├── MeetingError\n'
    '├── TranscriptionError\n'
    '├── AIProcessingError\n'
    '├── CalendarError\n'
    '├── ACSError\n'
    '├── DatabaseError\n'
    '├── CacheError\n'
    '├── GraphAPIError\n'
    '├── DeliveryError\n'
    '├── RateLimitError\n'
    '├── ValidationError\n'
    '└── ResourceNotFoundError'
)

doc.add_heading('18.2 HTTP Error Mapping', level=2)
add_table_simple(
    ['HTTP Code', 'When', 'Client Message'],
    [
        ('400', 'Validation failure, empty body', 'Specific validation error'),
        ('401', 'Missing/expired/invalid JWT', 'Token has expired / Invalid token'),
        ('403', 'Insufficient role permissions', 'Forbidden'),
        ('404', 'Entity not found', 'Meeting/Document/Action not found'),
        ('429', 'Webhook rate limit exceeded', 'Too many requests'),
        ('502', 'Upstream service failure', 'Generic message (no internal details)'),
    ],
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
# 19. OBSERVABILITY
# ══════════════════════════════════════════════════════════════════
doc.add_heading('19. Observability & Monitoring', level=1)

add_table_simple(
    ['Layer', 'Technology', 'Details'],
    [
        ('Logging', 'Structured JSON (python-json-logger)', 'Azure Monitor compatible, correlation IDs'),
        ('Tracing', 'OpenTelemetry', 'Spans, propagation, OTLP/Jaeger/Zipkin exporters'),
        ('Metrics', 'OpenTelemetry Metrics', 'Counters, histograms, gauges with no-op fallback'),
        ('Health Checks', 'GET /health per service', 'Docker healthcheck + k8s readiness probes'),
        ('M365 Status', 'GET /api/dashboard/m365-status', 'Checks Azure credential configuration'),
    ],
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
# 20. TESTING
# ══════════════════════════════════════════════════════════════════
doc.add_heading('20. Testing Strategy', level=1)

add_table_simple(
    ['Type', 'Framework', 'Count', 'Coverage'],
    [
        ('Backend Unit', 'pytest + pytest-asyncio', '323', 'All services + foundation'),
        ('Frontend Build', 'Angular CLI', 'n/a', 'TypeScript compilation + template validation'),
        ('Foundation Smoke', 'pytest', '143', 'Import tests + DSPy integration'),
        ('RAG Evaluation', 'RAGEvaluator', '20 cases', 'Golden QA: precision@k, MRR, faithfulness'),
        ('CI/CD', 'GitHub Actions', 'matrix', 'Per-service test jobs + lint + build'),
    ],
)

doc.add_heading('20.1 Test Execution', level=2)
add_code(
    '# All backend tests\n'
    'for svc in meeting document chat dashboard pre-meeting-brief weekly-digest; do\n'
    '  PYTHONPATH=services/${svc}-service/src:foundation/src pytest services/${svc}-service/tests/ -v\n'
    'done\n'
    'PYTHONPATH=foundation/src pytest foundation/tests/ -v\n'
    '\n'
    '# Frontend build\n'
    'cd yoda-frontend && npx ng build --configuration=production'
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
# 21. APPENDIX
# ══════════════════════════════════════════════════════════════════
doc.add_heading('21. Appendix: File Inventory', level=1)

add_table_simple(
    ['Area', 'Location', 'Files', 'Purpose'],
    [
        ('Foundation', 'YODA-BB/foundation/src/yoda_foundation/', '241', 'Shared library'),
        ('Meeting Service', 'YODA-BB/services/meeting-service/', '20', 'Meeting lifecycle'),
        ('Document Service', 'YODA-BB/services/document-service/', '12', 'Document management'),
        ('Chat Service', 'YODA-BB/services/chat-service/', '8', 'AI Q&A'),
        ('Dashboard Service', 'YODA-BB/services/dashboard-service/', '16', 'Stats & insights'),
        ('Brief Service', 'YODA-BB/services/pre-meeting-brief-service/', '8', 'Meeting prep'),
        ('Digest Service', 'YODA-BB/services/weekly-digest-service/', '8', 'Weekly summaries'),
        ('Frontend', 'yoda-frontend/src/app/', '40+', 'Angular SPA'),
        ('Deployment', 'YODA-BB/deployment/', '3', 'docker-compose, nginx, Dockerfiles'),
        ('Alembic', 'YODA-BB/alembic/', '3', 'DB migrations'),
        ('Config', 'YODA-BB/config/', '1', 'golden_qa.json'),
        ('CI/CD', 'YODA-BB/.github/workflows/', '1', 'ci.yml'),
        ('Docs', 'YODA-BB/docs/', '5+', 'Generated documentation'),
    ],
)

# ── Save ──────────────────────────────────────────────────────────
output_path = '/Users/srinivaasant/Documents/YODA/YODA-BB/docs/YODA_Low_Level_Design.docx'
doc.save(output_path)
print(f'Document saved to: {output_path}')
