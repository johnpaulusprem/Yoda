"""Generate YODA User Stories & Requirements Document (Word format).

Produces a comprehensive .docx covering all epics, user stories,
acceptance criteria, and technical notes for the YODA system.
"""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from datetime import datetime

doc = Document()

# ── Styles ────────────────────────────────────────────────────────
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)

for level in range(1, 4):
    hs = doc.styles[f'Heading {level}']
    hs.font.name = 'Calibri'
    hs.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)

# ── Title Page ────────────────────────────────────────────────────
doc.add_paragraph()
doc.add_paragraph()
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('YODA — AI Meeting Companion')
run.font.size = Pt(28)
run.font.color.rgb = RGBColor(0x3B, 0x82, 0xF6)
run.bold = True

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('User Stories & Requirements Specification')
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

doc.add_paragraph()
meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.add_run(f'Version 1.0 — {datetime.now().strftime("%B %d, %Y")}').font.size = Pt(12)

doc.add_paragraph()
meta2 = doc.add_paragraph()
meta2.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta2.add_run('Confidential — For Internal Use Only').font.italic = True

doc.add_page_break()

# ── Table of Contents placeholder ─────────────────────────────────
doc.add_heading('Table of Contents', level=1)
doc.add_paragraph('(Generate TOC in Word: References → Table of Contents → Automatic Table)')
doc.add_page_break()

# ── 1. Document Overview ──────────────────────────────────────────
doc.add_heading('1. Document Overview', level=1)
doc.add_paragraph(
    'This document defines the complete set of user stories for YODA (Your Organizational Digital Assistant), '
    'an AI-powered meeting companion for enterprise CXOs. YODA integrates with Microsoft 365 (Teams, SharePoint, '
    'OneDrive, Outlook, Calendar) to provide pre-meeting briefs, real-time transcription, post-meeting summaries, '
    'action item tracking, document intelligence, and AI-powered Q&A.'
)

doc.add_heading('1.1 Product Vision', level=2)
doc.add_paragraph(
    'Enable CXOs to walk into every meeting prepared, walk out with clear action items, '
    'and have an AI companion that remembers everything discussed across all meetings — '
    'accessible through a single enterprise-grade web application.'
)

doc.add_heading('1.2 User Personas', level=2)
table = doc.add_table(rows=4, cols=3, style='Light Grid Accent 1')
table.alignment = WD_TABLE_ALIGNMENT.CENTER
headers = ['Persona', 'Role', 'Key Needs']
for i, h in enumerate(headers):
    table.rows[0].cells[i].text = h
    table.rows[0].cells[i].paragraphs[0].runs[0].bold = True

personas = [
    ('Arun (CXO)', 'CEO / C-Suite Executive', 'Meeting prep, decision tracking, delegation visibility, weekly digest'),
    ('Priya (Direct Report)', 'VP / Senior Director', 'Action item status, document sharing, meeting summaries'),
    ('System Admin', 'IT Administrator', 'User management, bot configuration, compliance monitoring'),
]
for idx, (name, role, needs) in enumerate(personas, 1):
    table.rows[idx].cells[0].text = name
    table.rows[idx].cells[1].text = role
    table.rows[idx].cells[2].text = needs

doc.add_heading('1.3 System Architecture', level=2)
doc.add_paragraph(
    'YODA is built as a microservices backend (Python/FastAPI) with an Angular 21 frontend, '
    'deployed via Docker Compose behind an Nginx API gateway. The system comprises 6 backend services, '
    'a shared foundation library, PostgreSQL with pgvector for RAG, and Redis for caching.'
)

doc.add_page_break()

# ── Helper to add user story ──────────────────────────────────────
story_counter = [0]

def add_story(epic_id, title, as_a, i_want, so_that, acceptance_criteria, priority='Medium', api_endpoints=None, notes=None):
    story_counter[0] += 1
    sid = f'US-{story_counter[0]:03d}'

    doc.add_heading(f'{sid}: {title}', level=3)

    p = doc.add_paragraph()
    p.add_run('ID: ').bold = True
    p.add_run(sid)
    p.add_run('  |  Priority: ').bold = True
    p.add_run(priority)
    p.add_run('  |  Epic: ').bold = True
    p.add_run(epic_id)

    doc.add_paragraph(f'As a {as_a}, I want to {i_want}, so that {so_that}.')

    doc.add_paragraph('Acceptance Criteria:', style='List Bullet')
    for ac in acceptance_criteria:
        doc.add_paragraph(ac, style='List Bullet 2')

    if api_endpoints:
        p = doc.add_paragraph()
        p.add_run('API Endpoints: ').bold = True
        p.add_run(', '.join(api_endpoints))

    if notes:
        p = doc.add_paragraph()
        p.add_run('Technical Notes: ').bold = True
        p.add_run(notes)

    doc.add_paragraph()  # spacing


# ══════════════════════════════════════════════════════════════════
# EPIC 1: DASHBOARD
# ══════════════════════════════════════════════════════════════════
doc.add_heading('2. Epic 1: Executive Dashboard', level=1)
doc.add_paragraph(
    'The dashboard is the CXO\'s landing screen — a single-glance overview of today\'s meetings, '
    'pending actions, documents requiring review, and items needing immediate attention.'
)

add_story('E1-DASHBOARD', 'View daily KPI summary',
    'CXO', 'see my key metrics at a glance when I open the app',
    'I can quickly assess my day without clicking into multiple screens',
    [
        'Dashboard shows 4 stat cards: Meetings Today, Open Actions, Overdue Items, Docs to Review',
        'Each card shows the count and a trend indicator (up/down vs last week)',
        'Cards are clickable — navigate to the relevant detail view',
        'Data refreshes on each visit (no stale cache)',
    ],
    priority='High',
    api_endpoints=['GET /api/dashboard/stats'],
)

add_story('E1-DASHBOARD', 'View today\'s meeting schedule',
    'CXO', 'see all my meetings for today with key details',
    'I can plan my day and prepare for upcoming meetings',
    [
        'Shows list of today\'s meetings ordered by time',
        'Each meeting shows: time, subject, attendee count, location/channel',
        'Tags indicate status: Brief Ready, High Priority, Decision Needed, Prep Required',
        'Click a meeting opens its Pre-Meeting Brief',
        '"View all" link navigates to full Meetings view',
    ],
    priority='High',
    api_endpoints=['GET /api/meetings?limit=10'],
)

add_story('E1-DASHBOARD', 'Join meeting directly from dashboard',
    'CXO', 'join a Teams meeting with one click from the dashboard',
    'I don\'t have to switch to Teams or Calendar to join',
    [
        'Meetings within 5 minutes of start show a pulsing green "Join Now" button',
        'Meetings in progress show a blue "In Progress · Join" button',
        'Meetings starting within 1 hour show a yellow "Starts in Xm" indicator',
        'Clicking "Join Now" opens the Teams meeting in a new tab via the join_url',
        'Button state updates automatically based on current time vs scheduled_start',
    ],
    priority='High',
    api_endpoints=['GET /api/meetings (join_url + join_status fields)'],
    notes='join_status is a computed field on MeetingResponse based on current time and meeting status',
)

add_story('E1-DASHBOARD', 'View items needing attention',
    'CXO', 'see overdue and urgent items highlighted prominently',
    'I can take immediate action on critical items',
    [
        'Attention section shows overdue action items with days-overdue count',
        'Shows items due today with "DUE TODAY" tag',
        'Shows pending decisions with "Decision Pending" tag',
        'Items scoped to the authenticated user (tenant isolation)',
        'Color-coded by severity: red (overdue), yellow (due soon), purple (prep required)',
    ],
    priority='High',
    api_endpoints=['GET /api/dashboard/attention-items'],
)

add_story('E1-DASHBOARD', 'View recent activity feed',
    'CXO', 'see what happened recently across my meetings and documents',
    'I stay informed about team activity without checking each system',
    [
        'Shows timeline of: document updates, action completions, meeting summaries, mentions',
        'Each item shows icon, description, and relative timestamp',
        'Limited to 20 most recent items',
        'Scoped to user\'s meetings and documents',
    ],
    priority='Medium',
    api_endpoints=['GET /api/dashboard/activity-feed'],
)

add_story('E1-DASHBOARD', 'Use quick action shortcuts',
    'CXO', 'quickly navigate to key features from the dashboard',
    'I can access any feature in one click',
    [
        'Shows 5 shortcut buttons: Pre-Meeting Brief, Ask AI, Meeting Summary, Track Actions, Weekly Digest',
        'Each button navigates to the corresponding view',
    ],
    priority='Low',
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
# EPIC 2: MEETINGS
# ══════════════════════════════════════════════════════════════════
doc.add_heading('3. Epic 2: Meeting Management', level=1)
doc.add_paragraph(
    'Full meeting lifecycle management — from calendar sync through transcription, '
    'AI summarization, action item extraction, and delivery to Teams.'
)

add_story('E2-MEETINGS', 'View meetings grouped by day',
    'CXO', 'see my meetings organized by date with filter options',
    'I can plan ahead and review past meetings',
    [
        'Meetings grouped by day (Today, Tomorrow, date headers)',
        'Filter tabs: Today, This Week, This Month',
        'Each meeting shows: time, subject, attendee count, duration, tags',
        'Tags: Brief Ready, High Priority, Recurring, External',
        'Join button states match dashboard behavior',
    ],
    priority='High',
    api_endpoints=['GET /api/meetings?status=X&limit=20&offset=0'],
)

add_story('E2-MEETINGS', 'View post-meeting summary',
    'CXO', 'review the AI-generated summary after a meeting ends',
    'I can recall what was discussed without rewatching the recording',
    [
        'Shows: meeting subject, date, duration, participant count',
        'Key Discussion Points: structured paragraphs covering all topics',
        'Decisions Made: numbered list with context and conditions',
        'Action Items: table with description, owner, due date, priority, confidence',
        'Open Questions: items raised but not resolved',
        'Meeting Analytics: 4-stat grid (duration, participants, decisions, actions)',
    ],
    priority='High',
    api_endpoints=['GET /api/meetings/{id}'],
    notes='MeetingDetailResponse includes nested summary, action_items, and participants',
)

add_story('E2-MEETINGS', 'Edit meeting summary',
    'CXO', 'correct or amend the AI-generated summary',
    'I can fix inaccuracies before sharing with attendees',
    [
        'Edit button opens editable view of summary text, decisions, key topics',
        'Partial update — only changed fields are sent',
        'Returns 404 if meeting or summary not found',
        'Requires authentication',
    ],
    priority='Medium',
    api_endpoints=['PATCH /api/meetings/{id}/summary'],
)

add_story('E2-MEETINGS', 'Share summary to attendees',
    'CXO', 'send the meeting summary to all attendees via Teams',
    'everyone has a shared record of what was discussed and decided',
    [
        'Share button sends Adaptive Card to the Teams meeting chat',
        'Card includes summary, decisions, and action items',
        'Delivery status tracked (delivered_at timestamp)',
    ],
    priority='Medium',
    api_endpoints=['POST /api/meetings/{id}/reprocess'],
    notes='DeliveryService sends Adaptive Cards via Graph API',
)

add_story('E2-MEETINGS', 'Detect decision conflicts',
    'CXO', 'be alerted when a new decision contradicts a previous one',
    'I can avoid making inconsistent decisions across meetings',
    [
        'System compares new decisions against past 90 days of decisions',
        'Detects contradictions using keyword pair analysis (increase/decrease, approve/reject, etc.)',
        'Shows conflict alert with: current decision, past decision, meeting source, reason',
        'Conflict stored as MeetingInsight with severity=warning',
    ],
    priority='High',
    api_endpoints=['GET /api/insights (conflict_detection type)'],
    notes='ConflictDetectionService in dashboard-service compares decisions across meetings',
)

add_story('E2-MEETINGS', 'Automatic bot join and transcription',
    'CXO', 'have the bot automatically join my meetings and transcribe',
    'I get summaries without any manual action',
    [
        'CalendarWatcher monitors Graph webhooks for new/updated meetings',
        'Bot joins automatically when meeting starts (if user opted in)',
        'Transcript chunks arrive via webhook from the .NET media bot',
        'Chunks buffered and reassembled per speaker',
        'SSE endpoint streams live transcript to the frontend',
    ],
    priority='High',
    api_endpoints=['POST /api/bot-events/transcript', 'POST /api/bot-events/lifecycle', 'GET /api/meetings/{id}/events'],
    notes='SSE uses in-memory event queues per meeting. CalendarWatcher uses session_factory pattern.',
)

add_story('E2-MEETINGS', 'View meeting tags',
    'CXO', 'see computed tags on each meeting for quick scanning',
    'I can identify which meetings need attention at a glance',
    [
        'Tags computed: Brief Ready, High Priority, Has Actions, Recurring, External, Decision Needed, In Progress, Completed',
        'Recurring: detected from subject patterns (weekly, standup, 1:1, retro)',
        'External: any participant email domain differs from organizer',
        'Tags returned in MeetingWithTagsResponse',
    ],
    priority='Medium',
    api_endpoints=['GET /api/meetings (tags field in response)'],
    notes='MeetingTagService computes tags using batch-loaded action items and summaries (no N+1)',
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
# EPIC 3: PRE-MEETING BRIEF
# ══════════════════════════════════════════════════════════════════
doc.add_heading('4. Epic 3: Pre-Meeting Brief', level=1)
doc.add_paragraph(
    'AI-generated preparation material for each meeting — attendee context, '
    'past decisions, relevant documents, email threads, and suggested talking points.'
)

add_story('E3-BRIEF', 'View comprehensive pre-meeting brief',
    'CXO', 'see everything I need to know before walking into a meeting',
    'I am fully prepared without manual research',
    [
        'Brief shows: meeting title, time, location, duration, recurrence',
        'Countdown timer showing minutes until start',
        'Join Meeting button (large, green, pulsing when within 5 min)',
        'Attendee grid with: name, avatar, role, department, context (overdue items, last 1:1)',
        'Past decisions from the last meeting with this group',
        'Relevant documents updated this week (with file type icons)',
        'Related email threads (subject, sender, preview)',
        'AI-suggested topics and questions (5 clickable items)',
        'Click a suggested question navigates to Chat with the question pre-filled',
    ],
    priority='High',
    api_endpoints=['GET /api/briefs/{meeting_id}'],
    notes='PreMeetingService concurrently fetches: Graph (attendees), DB (past decisions), documents, emails, AI (questions)',
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
# EPIC 4: ACTION ITEMS
# ══════════════════════════════════════════════════════════════════
doc.add_heading('5. Epic 4: Action Item Tracking', level=1)
doc.add_paragraph(
    'Track, filter, and manage action items extracted from meetings. '
    'Includes nudge reminders, escalation, and delegation tracking.'
)

add_story('E4-ACTIONS', 'View action items by urgency',
    'CXO', 'see my action items organized by urgency level',
    'I can focus on the most critical items first',
    [
        'Four sections: Overdue (red), Due Soon (yellow), In Progress (blue), Upcoming (green)',
        'Each item shows: description, owner, due date, source meeting, status tag',
        'Overdue items show "X days overdue" count',
        'Source meeting is clickable — navigates to meeting detail',
    ],
    priority='High',
    api_endpoints=['GET /api/action-items?status=X&priority=Y'],
)

add_story('E4-ACTIONS', 'Filter action items',
    'CXO', 'filter actions by status, owner, and source meeting',
    'I can find specific items quickly',
    [
        'Status filter: All, Open, Overdue, In Progress, Completed',
        'Owner filter: All, Assigned to Me, Assigned by Me, My Direct Reports',
        'Source meeting filter: dynamically populated from meetings with actions',
        '"My Direct Reports" resolves via Graph API /users/{id}/directReports with caching',
    ],
    priority='High',
    api_endpoints=['GET /api/action-items?filter=assigned_to_me|assigned_by_me|my_reports'],
)

add_story('E4-ACTIONS', 'Complete and snooze action items',
    'CXO', 'mark items as complete or snooze reminders',
    'I can manage my commitments efficiently',
    [
        'Complete button: sets status=completed and completed_at timestamp',
        'Snooze button: pauses nudge reminders for N days',
        'Update action: change status, priority, deadline, or assignee',
        'Completed items removed from active view',
    ],
    priority='High',
    api_endpoints=['POST /api/action-items/{id}/complete', 'POST /api/action-items/{id}/snooze', 'PATCH /api/action-items/{id}'],
)

add_story('E4-ACTIONS', 'Receive nudge reminders',
    'CXO', 'get reminded about overdue or upcoming action items',
    'nothing falls through the cracks',
    [
        'NudgeScheduler runs periodically (configurable interval)',
        'Sends Adaptive Card reminders to Teams for overdue items',
        'Respects snooze: snoozed items skipped until snooze_until expires',
        'Escalation: after configurable threshold, escalates to manager',
        'Cooldown: doesn\'t re-nudge within cooldown period',
    ],
    priority='Medium',
    api_endpoints=['Background task — no API endpoint'],
    notes='NudgeScheduler uses session_factory pattern for per-run DB sessions',
)

add_story('E4-ACTIONS', 'AI confidence on action items',
    'CXO', 'see how confident the AI is about each extracted action item',
    'I can validate and correct low-confidence items',
    [
        'Each action item has a confidence score: high (0.9), medium (0.6), low (0.3)',
        'Extracted by the AI processor alongside priority',
        'Displayed as a badge on the action item row',
    ],
    priority='Low',
    notes='AI prompt includes "confidence: high|medium|low" field, mapped to float on ActionItem model',
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
# EPIC 5: DOCUMENTS
# ══════════════════════════════════════════════════════════════════
doc.add_heading('6. Epic 5: Document Intelligence', level=1)
doc.add_paragraph(
    'SharePoint/OneDrive document sync, AI classification, semantic search, '
    'and integration with the RAG pipeline for AI-powered Q&A.'
)

add_story('E5-DOCS', 'View and filter documents',
    'CXO', 'see my documents organized by type and recency',
    'I can find what I need without searching through folders',
    [
        'Type filter: Presentations, Spreadsheets, Documents, PDFs',
        'Sort options: Recently Updated, Most Relevant, Shared with Me',
        'Each document shows: title, type icon, updated by, folder path, file size',
        'Documents scoped to user (uploaded_by or shared_by)',
    ],
    priority='High',
    api_endpoints=['GET /api/documents?doc_type=X&sort=Y', 'GET /api/documents/recent'],
)

add_story('E5-DOCS', 'View documents needing review',
    'CXO', 'see documents shared with me that need my review or approval',
    'I don\'t miss important documents waiting for my input',
    [
        'Needs Review section shows documents with review_status=pending_review or action_required',
        'Priority tags: High Priority, Action Required, Pending Review',
        'Ordered by priority (high first), then shared_at descending',
    ],
    priority='High',
    api_endpoints=['GET /api/documents/needs-review'],
)

add_story('E5-DOCS', 'Sync documents from SharePoint/OneDrive',
    'CXO', 'pull my recent documents from Microsoft 365',
    'the system always has my latest files indexed',
    [
        '"Sync from M365" button triggers Graph API pull',
        'Fetches recent OneDrive items via /users/{id}/drive/recent',
        'Deduplicates by source_url or graph_item_id',
        'New documents auto-classified by the DocumentClassifier (23 categories)',
        'Sets folder_path, page_count, last_modified_by from Graph metadata',
    ],
    priority='Medium',
    api_endpoints=['POST /api/documents/sync'],
)

add_story('E5-DOCS', 'Semantic search across documents',
    'CXO', 'search my documents using natural language',
    'I can find information even when I don\'t remember exact keywords',
    [
        'Search box accepts natural language queries',
        'Uses hybrid search: vector similarity (pgvector) + full-text (tsvector)',
        'Results merged via Reciprocal Rank Fusion (RRF)',
        'LLM re-ranker re-scores top 15 candidates, returns best 5',
        'HyDE query expansion generates hypothetical answer for better retrieval',
        'Metadata-aware: auto-detects time/type/meeting filters from query text',
    ],
    priority='High',
    api_endpoints=['GET /api/documents/search?q=X'],
    notes='RAG pipeline: hybrid_retriever → reranker → context_builder → ChainOfThought',
)

add_story('E5-DOCS', 'Auto-classify documents',
    'CXO', 'have documents automatically categorized and tagged',
    'I can find documents by category without manual tagging',
    [
        '23 enterprise categories: MBR, QBR, SOW, MSA, Status Report, Delivery, Risk, MOM, Escalation, etc.',
        'Classification via vector template matching with keyword boosting',
        'File-format-aware: uses filename patterns + MIME type signals',
        'Assigns: category, confidence score, suggested priority, suggested tags',
        'Auto-classifies during process_document() and on-demand via API',
    ],
    priority='Medium',
    api_endpoints=['POST /api/documents/{id}/classify', 'POST /api/documents/classify-text'],
)

add_story('E5-DOCS', 'Index emails for AI search',
    'CXO', 'have my recent emails searchable through the AI chat',
    'the AI can answer questions from both documents and emails',
    [
        'Fetches recent emails from Graph API /users/{id}/messages',
        'Strips HTML, creates Document records with source=email',
        'Runs through ingestion pipeline (chunk + embed into pgvector)',
        'Available for RAG retrieval in Chat/Ask AI',
    ],
    priority='Medium',
    api_endpoints=['POST /api/documents/index-emails?days=7'],
)

add_story('E5-DOCS', 'View meeting-related documents',
    'CXO', 'see which documents are relevant to today\'s meetings',
    'I can review materials before each meeting',
    [
        'Cross-references today\'s calendar events with local documents and Graph attachments',
        'Shows document title, type, and which meeting it relates to',
    ],
    priority='Medium',
    api_endpoints=['GET /api/documents/meeting-related'],
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
# EPIC 6: CHAT / ASK AI
# ══════════════════════════════════════════════════════════════════
doc.add_heading('7. Epic 6: AI-Powered Chat (Ask AI)', level=1)
doc.add_paragraph(
    'RAG-powered conversational AI that answers questions from meetings, documents, '
    'calendar, and emails — with source citations for every answer.'
)

add_story('E6-CHAT', 'Ask questions with cited answers',
    'CXO', 'ask natural language questions and get AI answers with sources',
    'I can get instant answers without searching through multiple systems',
    [
        'User sends message → RAG pipeline retrieves relevant chunks → AI generates answer',
        'Every answer includes source citations (document name, meeting, URL)',
        'Citations are clickable — open the source document or meeting',
        'Chain-of-Thought reasoning with DSPy for higher quality answers',
        'dspy.Suggest assertions enforce: answer must cite sources, answer must be substantive',
    ],
    priority='High',
    api_endpoints=['POST /api/chat/sessions/{id}/messages'],
    notes='RAGPipeline: HyDE expansion → hybrid search → LLM rerank → context build → ChainOfThought → citation resolve',
)

add_story('E6-CHAT', 'Manage chat sessions',
    'CXO', 'create and manage multiple chat conversations',
    'I can organize different topics into separate threads',
    [
        'Create new session with a title',
        'List all sessions ordered by last activity',
        'View full message history for a session',
        'Empty state shows suggested prompts',
    ],
    priority='Medium',
    api_endpoints=['POST /api/chat/sessions', 'GET /api/chat/sessions', 'GET /api/chat/sessions/{id}/messages'],
)

add_story('E6-CHAT', 'Use suggested prompts',
    'CXO', 'see helpful prompt suggestions to start a conversation',
    'I know what kinds of questions the AI can answer',
    [
        'Empty state shows 5 suggestions: "What needs my attention today?", "What has Ravi committed to this month?", etc.',
        'Pre-meeting brief suggested questions are clickable → navigate to Chat with pre-filled query',
        'Clicking a suggestion sends it as a message immediately',
    ],
    priority='Low',
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
# EPIC 7: INSIGHTS
# ══════════════════════════════════════════════════════════════════
doc.add_heading('8. Epic 7: AI-Generated Insights', level=1)
doc.add_paragraph(
    'AI-detected patterns, trends, and anomalies from meetings and actions — '
    'helping the CXO understand time allocation, team performance, and decision patterns.'
)

add_story('E7-INSIGHTS', 'View meeting time analysis',
    'CXO', 'see how much time I spend in meetings',
    'I can optimize my calendar and delegate where appropriate',
    [
        'Shows total meetings and average hours per week over configurable period',
        'Period selector: 7, 30, 90, 365 days',
    ],
    priority='Medium',
    api_endpoints=['GET /api/insights/meeting-time?days=30'],
)

add_story('E7-INSIGHTS', 'View action item completion rate',
    'CXO', 'see the completion rate across my organization',
    'I can identify bottlenecks and underperforming areas',
    [
        'Shows total items, completed count, completion percentage',
        'Configurable time period',
    ],
    priority='Medium',
    api_endpoints=['GET /api/insights/action-completion?days=30'],
)

add_story('E7-INSIGHTS', 'View decision velocity',
    'CXO', 'see how quickly decisions are being resolved',
    'I can identify decision bottlenecks',
    [
        'Shows average days from decision-type action item creation to completion',
        'Identifies decision items by keywords: approve, decide, sign off, review, finalize, authorize',
    ],
    priority='Medium',
    api_endpoints=['GET /api/insights/decision-velocity?days=30'],
)

add_story('E7-INSIGHTS', 'View collaboration patterns',
    'CXO', 'see who I interact with most and identify stale relationships',
    'I can maintain key relationships and schedule check-ins',
    [
        'Top collaborators: ranked by meeting frequency',
        'Stale 1:1s: contacts not met in 14+ days',
    ],
    priority='Medium',
    api_endpoints=['GET /api/insights/collaboration?days=30'],
)

add_story('E7-INSIGHTS', 'Detect recurring topics',
    'CXO', 'be alerted when the same topic comes up in 3+ meetings',
    'I can schedule a dedicated session to resolve persistent issues',
    [
        'Analyzes key_topics from meeting summaries over 30 days',
        'Groups by semantic similarity',
        'Topics in 3+ unique meetings flagged as recurring',
        'Sorted by meeting count descending',
    ],
    priority='Medium',
    api_endpoints=['GET /api/insights/recurring-topics'],
    notes='RecurringTopicService in dashboard-service',
)

add_story('E7-INSIGHTS', 'Get AI recommendations',
    'CXO', 'receive proactive recommendations based on my data',
    'I can take action before issues escalate',
    [
        'Rule-based recommendations: overdue clusters (3+ items per person), stale 1:1s, unresolved recurring topics',
        'Each recommendation has: type, title, description, priority',
    ],
    priority='Low',
    api_endpoints=['GET /api/dashboard/recommendations'],
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
# EPIC 8: WEEKLY DIGEST
# ══════════════════════════════════════════════════════════════════
doc.add_heading('9. Epic 8: Weekly Digest', level=1)

add_story('E8-DIGEST', 'View weekly digest',
    'CXO', 'see a comprehensive summary of my week',
    'I can review the week\'s activity in one place',
    [
        'Stats grid: total meetings, time in meetings, decisions made, actions created',
        'Key Decisions This Week: list with source meetings',
        'Items Needing Follow-up: overdue and at-risk items',
        'Project Updates: status and owner per project',
        'People Notes: stale 1:1s, milestone reminders',
    ],
    priority='Medium',
    api_endpoints=['GET /api/digests/latest'],
)

add_story('E8-DIGEST', 'Auto-generate digest on schedule',
    'CXO', 'receive the weekly digest automatically every Friday',
    'I don\'t have to remember to check it',
    [
        'APScheduler generates digest on configurable day/time (default: Friday 4 PM UTC)',
        'Configurable user list (DIGEST_USER_IDS env var)',
        'Manual generation also available via API',
    ],
    priority='Medium',
    api_endpoints=['POST /api/digests/generate'],
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
# EPIC 9: NOTIFICATIONS & SEARCH
# ══════════════════════════════════════════════════════════════════
doc.add_heading('10. Epic 9: Notifications & Global Search', level=1)

add_story('E9-NOTIF', 'View and manage notifications',
    'CXO', 'see and manage my notifications in the app',
    'I stay informed about important events',
    [
        'Bell icon in topbar shows unread count badge',
        'Dropdown lists recent notifications with read/unread state',
        'Mark individual notification as read',
        'Mark all as read with one click',
        'Click notification navigates to related meeting or action',
    ],
    priority='Medium',
    api_endpoints=['GET /api/notifications', 'GET /api/notifications/count', 'PATCH /api/notifications/{id}/read', 'POST /api/notifications/read-all'],
)

add_story('E9-NOTIF', 'Global search across entities',
    'CXO', 'search across meetings, documents, and action items from the topbar',
    'I can find anything in the system quickly',
    [
        'Search box in topbar with live dropdown results',
        'Results grouped by type: meetings, documents, action items',
        'Each result shows: type icon, title, snippet, type badge',
        'Click result navigates to the relevant detail view',
    ],
    priority='Medium',
    api_endpoints=['GET /api/search?q=X&types=meetings,documents,actions'],
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
# EPIC 10: SETTINGS & ADMIN
# ══════════════════════════════════════════════════════════════════
doc.add_heading('11. Epic 10: Settings & Administration', level=1)

add_story('E10-SETTINGS', 'Manage user preferences',
    'CXO', 'configure my notification and digest preferences',
    'I receive information in my preferred format and schedule',
    [
        'Theme toggle: dark/light mode (persisted to localStorage)',
        'Summary delivery: chat, email, or both',
        'Action item nudges: enable/disable',
        'Weekly digest: enable/disable',
        'Auto-join meetings: enable/disable',
        'Preferences persisted to UserPreference model via API',
    ],
    priority='Medium',
    api_endpoints=['GET /api/settings', 'PATCH /api/settings'],
)

add_story('E10-SETTINGS', 'View M365 connection status',
    'CXO', 'see whether the system is connected to Microsoft 365',
    'I know if my data is up to date',
    [
        'Status indicator in topbar: green dot + "Connected to M365"',
        'Health endpoint checks if Azure credentials are configured',
    ],
    priority='Low',
    api_endpoints=['GET /api/dashboard/m365-status'],
)

add_story('E10-SETTINGS', 'Admin: manage users',
    'System Admin', 'create, update, and delete user accounts',
    'I can control who has access to the system',
    [
        'List users with optional opted_in filter',
        'Create user with preferences',
        'Update user preferences',
        'Delete user',
        'Admin-only endpoints (require Admin role)',
    ],
    priority='Medium',
    api_endpoints=['GET /api/admin/users', 'POST /api/admin/users', 'PATCH /api/admin/users/{id}', 'DELETE /api/admin/users/{id}'],
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
# EPIC 11: CROSS-CUTTING / NON-FUNCTIONAL
# ══════════════════════════════════════════════════════════════════
doc.add_heading('12. Epic 11: Cross-Cutting & Non-Functional Requirements', level=1)

add_story('E11-NFR', 'Azure AD SSO authentication',
    'CXO', 'sign in with my Microsoft 365 account',
    'I don\'t need a separate password for this system',
    [
        'MSAL.js handles login redirect to Azure AD / Entra ID',
        'JWT access token validated by backend (RS256, JWKS rotation)',
        'Roles: CXO.Admin, CXO.User, CXO.Viewer — mapped to RBAC permissions',
        'Dev mode: REQUIRE_AUTH=false bypasses authentication',
        'Token attached to all /api/* requests via Angular HTTP interceptor',
    ],
    priority='High',
)

add_story('E11-NFR', 'Tenant isolation',
    'CXO', 'only see data from my organization',
    'my data is secure and private',
    [
        'All query endpoints scope results by authenticated user_id',
        'Meetings: user is organizer or participant',
        'Action items: user is assignee or organized the meeting',
        'Documents: user is uploader or recipient of share',
        'Dashboard stats: scoped to user\'s meetings and items',
    ],
    priority='High',
)

add_story('E11-NFR', 'Security hardening',
    'System Admin', 'ensure the system is protected against common attacks',
    'we meet enterprise security requirements',
    [
        'Webhook clientState validation (prevents forgery)',
        'Rate limiting on webhook endpoint (100 req/min/IP)',
        'CORS restricted to specific methods and headers',
        'Error messages scrubbed — no internal details leaked to client',
        'SQL injection prevented: ORM parameterized queries + frozen PGVectorConfig',
        'XSS prevented: Angular sanitization, zero innerHTML',
        'Async JWKS fetch (no event loop blocking)',
        'lazy="raise" on all model relationships (prevents MissingGreenlet)',
    ],
    priority='High',
)

add_story('E11-NFR', 'Real-time meeting updates via SSE',
    'CXO', 'see live transcript updates during an active meeting',
    'I can follow along without being in the Teams call',
    [
        'SSE endpoint: GET /api/meetings/{id}/events',
        'Streams transcript chunks as they arrive from the bot',
        'Keepalive heartbeat every 30 seconds',
        'Auto-cleanup on client disconnect',
    ],
    priority='Low',
    api_endpoints=['GET /api/meetings/{id}/events'],
)

add_story('E11-NFR', 'RAG pipeline with DSPy optimization',
    'System Admin', 'ensure the AI answers are accurate and improvable',
    'we can tune the system\'s accuracy over time',
    [
        'Real DSPy v3.1.3 integration with BootstrapFewShot optimizer',
        'Golden QA cases stored in config/golden_qa.json (20 cases, 8 categories)',
        'dspy.Suggest assertions: answer must cite sources, must be substantive',
        'Evaluation framework: precision@k, recall@k, MRR, LLM-judged faithfulness',
        'Optimization from JSON file — no code changes needed to add cases',
    ],
    priority='Medium',
    notes='Run optimizer: YodaOptimizer(YodaQA()).optimize_from_file("config/golden_qa.json")',
)

doc.add_page_break()

# ── Summary Table ─────────────────────────────────────────────────
doc.add_heading('13. Story Summary', level=1)

table = doc.add_table(rows=1, cols=4, style='Light Grid Accent 1')
table.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(['Epic', 'Stories', 'High Priority', 'Total AC']):
    table.rows[0].cells[i].text = h
    table.rows[0].cells[i].paragraphs[0].runs[0].bold = True

epics = {
    'E1-DASHBOARD': 'Executive Dashboard',
    'E2-MEETINGS': 'Meeting Management',
    'E3-BRIEF': 'Pre-Meeting Brief',
    'E4-ACTIONS': 'Action Item Tracking',
    'E5-DOCS': 'Document Intelligence',
    'E6-CHAT': 'AI-Powered Chat',
    'E7-INSIGHTS': 'AI Insights',
    'E8-DIGEST': 'Weekly Digest',
    'E9-NOTIF': 'Notifications & Search',
    'E10-SETTINGS': 'Settings & Admin',
    'E11-NFR': 'Cross-Cutting / NFR',
}
# This is a placeholder — actual counts would be computed
for epic_id, name in epics.items():
    row = table.add_row()
    row.cells[0].text = name
    row.cells[1].text = '-'
    row.cells[2].text = '-'
    row.cells[3].text = '-'

doc.add_paragraph()
doc.add_paragraph(f'Total User Stories: {story_counter[0]}')

# ── Save ──────────────────────────────────────────────────────────
output_path = '/Users/srinivaasant/Documents/YODA/YODA-BB/docs/YODA_User_Stories_Requirements.docx'
doc.save(output_path)
print(f'Document saved to: {output_path}')
print(f'Total user stories: {story_counter[0]}')
