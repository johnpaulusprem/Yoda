"""Generate YODA Database Schema Document (Word format).

Produces a .docx with all table definitions, columns, types,
relationships, indexes, and ER description.
"""

import sys
sys.path.insert(0, 'foundation/src')

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from datetime import datetime

from yoda_foundation.models.base import Base
from yoda_foundation.models.meeting import Meeting, MeetingParticipant
from yoda_foundation.models.transcript import TranscriptSegment
from yoda_foundation.models.summary import MeetingSummary
from yoda_foundation.models.action_item import ActionItem
from yoda_foundation.models.subscription import GraphSubscription, UserPreference
from yoda_foundation.models.document import Document as DocModel, DocumentChunk
from yoda_foundation.models.insight import MeetingInsight, WeeklyDigest
from yoda_foundation.models.chat import ChatSession, ChatMessage
from yoda_foundation.models.notification import Notification
from yoda_foundation.models.project import Project

doc = Document()

# Styles
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(10)

for level in range(1, 4):
    hs = doc.styles[f'Heading {level}']
    hs.font.name = 'Calibri'
    hs.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)

# Title
doc.add_paragraph()
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('YODA — Database Schema Reference')
run.font.size = Pt(24)
run.font.color.rgb = RGBColor(0x3B, 0x82, 0xF6)
run.bold = True

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run(f'Version 1.0 — {datetime.now().strftime("%B %d, %Y")}')
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

doc.add_page_break()

# Overview
doc.add_heading('1. Overview', level=1)
doc.add_paragraph(
    'YODA uses a single PostgreSQL database with the pgvector extension for vector similarity search. '
    'All 6 microservices share this database, with Alembic managing schema migrations centrally. '
    'All tables use UUID primary keys and created_at/updated_at timestamp columns.'
)

doc.add_heading('1.1 Database Configuration', level=2)
p = doc.add_paragraph()
p.add_run('Engine: ').bold = True
p.add_run('PostgreSQL 16 with pgvector extension')
p = doc.add_paragraph()
p.add_run('Driver: ').bold = True
p.add_run('asyncpg (async)')
p = doc.add_paragraph()
p.add_run('ORM: ').bold = True
p.add_run('SQLAlchemy 2.0 async with Mapped type hints')
p = doc.add_paragraph()
p.add_run('Migrations: ').bold = True
p.add_run('Alembic (centralized at YODA-BB/alembic/)')

doc.add_heading('1.2 Table Summary', level=2)

# Table descriptions
table_descriptions = {
    'meetings': 'Core table — every Teams meeting tracked by the system. Source of truth for meeting lifecycle.',
    'meeting_participants': 'Attendees who joined each meeting. Populated from Graph API and bot events.',
    'transcript_segments': 'Speaker-attributed transcript chunks from the .NET media bot. Ordered by sequence_number.',
    'meeting_summaries': 'AI-generated summaries with decisions, key topics, and unresolved questions. One per meeting.',
    'action_items': 'Tasks extracted from meetings by the AI processor. Tracked through completion with nudge reminders.',
    'documents': 'Files from SharePoint/OneDrive/email/uploads. Processed through RAG pipeline for AI search.',
    'document_chunks': 'Text chunks with 1536-dim vector embeddings for pgvector similarity search.',
    'chat_sessions': 'User chat conversations with the AI assistant.',
    'chat_messages': 'Individual messages within chat sessions, including AI responses with source citations.',
    'notifications': 'System notifications (summary delivered, nudge sent, etc.).',
    'meeting_insights': 'AI-detected insights: conflict detection, sentiment, participation patterns, topic trends.',
    'weekly_digests': 'Auto-generated weekly summaries with stats, decisions, and follow-ups.',
    'graph_subscriptions': 'Active Microsoft Graph webhook subscriptions for calendar change notifications.',
    'user_preferences': 'Per-user settings: notification channel, auto-join, nudge, digest preferences.',
    'projects': 'Project entities linked to meetings for tracking cross-meeting initiatives.',
    'project_meetings': 'Many-to-many association between projects and meetings.',
}

summary_table = doc.add_table(rows=1, cols=3, style='Light Grid Accent 1')
summary_table.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(['Table', 'Rows (est.)', 'Description']):
    summary_table.rows[0].cells[i].text = h
    summary_table.rows[0].cells[i].paragraphs[0].runs[0].bold = True

for table_name in sorted(Base.metadata.tables.keys()):
    row = summary_table.add_row()
    row.cells[0].text = table_name
    row.cells[1].text = '-'
    row.cells[2].text = table_descriptions.get(table_name, '')

doc.add_page_break()

# Detailed schema for each table
doc.add_heading('2. Table Definitions', level=1)

for table_name, table in sorted(Base.metadata.tables.items()):
    doc.add_heading(f'2.x {table_name}', level=2)

    desc = table_descriptions.get(table_name, '')
    if desc:
        doc.add_paragraph(desc)

    # Column table
    col_table = doc.add_table(rows=1, cols=6, style='Light Grid Accent 1')
    col_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ['Column', 'Type', 'Nullable', 'PK', 'FK', 'Default/Notes']
    for i, h in enumerate(headers):
        col_table.rows[0].cells[i].text = h
        col_table.rows[0].cells[i].paragraphs[0].runs[0].bold = True
        col_table.rows[0].cells[i].paragraphs[0].runs[0].font.size = Pt(9)

    for col in table.columns:
        row = col_table.add_row()
        row.cells[0].text = col.name
        row.cells[0].paragraphs[0].runs[0].font.size = Pt(9)

        type_str = str(col.type)
        if 'UUID' in type_str:
            type_str = 'UUID'
        elif 'VECTOR' in type_str:
            type_str = 'VECTOR(1536)'
        row.cells[1].text = type_str
        row.cells[1].paragraphs[0].runs[0].font.size = Pt(9)

        row.cells[2].text = 'YES' if col.nullable else 'NO'
        row.cells[2].paragraphs[0].runs[0].font.size = Pt(9)

        row.cells[3].text = 'PK' if col.primary_key else ''
        row.cells[3].paragraphs[0].runs[0].font.size = Pt(9)

        fk_str = ''
        if col.foreign_keys:
            fk_str = list(col.foreign_keys)[0].target_fullname
        row.cells[4].text = fk_str
        row.cells[4].paragraphs[0].runs[0].font.size = Pt(9)

        notes = []
        if col.unique:
            notes.append('UNIQUE')
        if col.default and hasattr(col.default, 'arg') and col.default.arg is not None:
            val = col.default.arg
            if callable(val):
                notes.append('AUTO (uuid4/now)')
            else:
                notes.append(f'DEFAULT: {val}')
        if col.index:
            notes.append('INDEXED')
        row.cells[5].text = ', '.join(notes) if notes else ''
        row.cells[5].paragraphs[0].runs[0].font.size = Pt(9)

    doc.add_paragraph()

doc.add_page_break()

# Relationships
doc.add_heading('3. Relationships (Entity-Relationship)', level=1)

relationships = [
    ('meetings', 'meeting_participants', '1:N', 'meeting_id', 'Participants who joined'),
    ('meetings', 'transcript_segments', '1:N', 'meeting_id', 'Transcript chunks from bot'),
    ('meetings', 'meeting_summaries', '1:1', 'meeting_id (UNIQUE)', 'AI-generated summary'),
    ('meetings', 'action_items', '1:N', 'meeting_id', 'Extracted action items'),
    ('meetings', 'documents', '1:N', 'meeting_id', 'Related documents'),
    ('meetings', 'meeting_insights', '1:N', 'meeting_id', 'AI insights (conflicts, patterns)'),
    ('meetings', 'project_meetings', 'N:M', 'project_meetings join table', 'Projects linked to meetings'),
    ('documents', 'document_chunks', '1:N', 'document_id (CASCADE)', 'Text chunks with embeddings'),
    ('chat_sessions', 'chat_messages', '1:N', 'session_id', 'Messages in a chat thread'),
]

rel_table = doc.add_table(rows=1, cols=5, style='Light Grid Accent 1')
rel_table.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(['Parent', 'Child', 'Cardinality', 'FK Column', 'Description']):
    rel_table.rows[0].cells[i].text = h
    rel_table.rows[0].cells[i].paragraphs[0].runs[0].bold = True

for parent, child, card, fk, desc in relationships:
    row = rel_table.add_row()
    row.cells[0].text = parent
    row.cells[1].text = child
    row.cells[2].text = card
    row.cells[3].text = fk
    row.cells[4].text = desc

doc.add_page_break()

# Indexes
doc.add_heading('4. Indexes', level=1)
doc.add_paragraph('Beyond primary keys and foreign key indexes, the following explicit indexes exist:')

indexes = [
    ('chat_sessions', 'ix_chat_sessions_user_id', 'user_id', 'List sessions by user'),
    ('notifications', 'ix_notifications_user_id', 'user_id', 'List notifications by user'),
    ('weekly_digests', 'ix_weekly_digests_user_id', 'user_id', 'List digests by user'),
    ('document_chunks', 'ix_document_chunks_document_id', 'document_id', 'Chunks for a document'),
    ('documents', 'uq_documents_graph_item_id', 'graph_item_id (UNIQUE)', 'Dedup SharePoint items'),
    ('document_chunks', 'uq_document_chunk_index', '(document_id, chunk_index) UNIQUE', 'No duplicate chunks'),
    ('graph_subscriptions', 'uq_graph_sub_id', 'subscription_id (UNIQUE)', 'Dedup subscriptions'),
    ('user_preferences', 'uq_user_prefs_user_id', 'user_id (UNIQUE)', 'One preference per user'),
    ('meeting_summaries', 'uq_summaries_meeting_id', 'meeting_id (UNIQUE)', 'One summary per meeting'),
]

idx_table = doc.add_table(rows=1, cols=4, style='Light Grid Accent 1')
idx_table.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(['Table', 'Index Name', 'Column(s)', 'Purpose']):
    idx_table.rows[0].cells[i].text = h
    idx_table.rows[0].cells[i].paragraphs[0].runs[0].bold = True

for tbl, name, cols, purpose in indexes:
    row = idx_table.add_row()
    row.cells[0].text = tbl
    row.cells[1].text = name
    row.cells[2].text = cols
    row.cells[3].text = purpose

doc.add_paragraph()
doc.add_paragraph(
    'Note: pgvector automatically creates an index on VECTOR columns for cosine similarity search. '
    'For large datasets (>100K chunks), consider adding an IVFFlat or HNSW index on document_chunks.embedding.'
)

doc.add_page_break()

# Migration info
doc.add_heading('5. Migration Management', level=1)
doc.add_paragraph(
    'Migrations are managed centrally using Alembic at YODA-BB/alembic/. '
    'All model imports are in alembic/env.py which imports from yoda_foundation.models.*.'
)

doc.add_heading('5.1 Commands', level=2)
commands = [
    ('Generate migration', 'alembic revision --autogenerate -m "description"'),
    ('Apply all migrations', 'alembic upgrade head'),
    ('Rollback one step', 'alembic downgrade -1'),
    ('View current version', 'alembic current'),
    ('View migration history', 'alembic history'),
]
for desc, cmd in commands:
    p = doc.add_paragraph()
    p.add_run(f'{desc}: ').bold = True
    p.add_run(cmd).font.name = 'Consolas'

doc.add_heading('5.2 Environment', level=2)
doc.add_paragraph('Required: DATABASE_URL=postgresql+asyncpg://yoda:yoda_dev@localhost:5432/yoda')
doc.add_paragraph('Required: PYTHONPATH=foundation/src (so alembic can import yoda_foundation.models)')

# Scripts
doc.add_page_break()
doc.add_heading('6. Operational Scripts', level=1)

scripts = [
    ('Start infrastructure', 'cd deployment && docker compose up -d postgres redis', 'Starts PostgreSQL + Redis in Docker'),
    ('Create tables', 'DATABASE_URL=... alembic upgrade head', 'Applies all pending migrations'),
    ('Start meeting-service', 'PYTHONPATH=services/meeting-service/src:foundation/src uvicorn meeting_service.main:app --port 8010', 'Starts on port 8010'),
    ('Start document-service', 'PYTHONPATH=services/document-service/src:foundation/src uvicorn document_service.main:app --port 8011', ''),
    ('Start chat-service', 'PYTHONPATH=services/chat-service/src:foundation/src uvicorn chat_service.main:app --port 8012', ''),
    ('Start dashboard-service', 'PYTHONPATH=services/dashboard-service/src:foundation/src uvicorn dashboard_service.main:app --port 8013', ''),
    ('Start pre-meeting-brief', 'PYTHONPATH=services/pre-meeting-brief-service/src:foundation/src uvicorn pre_meeting_brief_service.main:app --port 8014', ''),
    ('Start weekly-digest', 'PYTHONPATH=services/weekly-digest-service/src:foundation/src uvicorn weekly_digest_service.main:app --port 8015', ''),
    ('Start Angular frontend', 'cd yoda-frontend && npx ng serve --port 4210', 'Dev server with API proxy'),
    ('Start everything (Docker)', 'cd deployment && docker compose up', 'All 10 containers'),
    ('Run all backend tests', 'for svc in meeting document chat dashboard pre-meeting-brief weekly-digest; do PYTHONPATH=... pytest ...; done', '323 tests'),
    ('Build frontend', 'cd yoda-frontend && npx ng build --configuration=production', '520KB output'),
    ('Optimize DSPy prompts', 'from yoda_foundation.dspy.integration import ...; optimizer.optimize_from_file("config/golden_qa.json")', 'Uses 20 golden QA cases'),
]

script_table = doc.add_table(rows=1, cols=3, style='Light Grid Accent 1')
script_table.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(['Task', 'Command', 'Notes']):
    script_table.rows[0].cells[i].text = h
    script_table.rows[0].cells[i].paragraphs[0].runs[0].bold = True

for task, cmd, notes in scripts:
    row = script_table.add_row()
    row.cells[0].text = task
    row.cells[1].text = cmd
    row.cells[1].paragraphs[0].runs[0].font.size = Pt(8)
    row.cells[1].paragraphs[0].runs[0].font.name = 'Consolas'
    row.cells[2].text = notes

# Save
output_path = '/Users/srinivaasant/Documents/YODA/YODA-BB/docs/YODA_Database_Schema_Reference.docx'
doc.save(output_path)
print(f'Document saved to: {output_path}')
print(f'Tables documented: {len(Base.metadata.tables)}')
